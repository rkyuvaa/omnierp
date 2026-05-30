import docx
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
import os

def create_document():
    doc = Document()

    # --- STYLE SETTING CONFIGURATION ---
    # Set page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # UI Branding Colors (Konwert UI Green Theme)
    COLOR_PRIMARY = RGBColor(25, 84, 2)     # #195402 - Deep UI Green
    COLOR_SECONDARY = RGBColor(35, 130, 4)   # #238204 - Lighter Green
    COLOR_TEXT = RGBColor(45, 55, 72)        # #2D3748 - Slate/Charcoal Body
    COLOR_MUTED = RGBColor(113, 128, 150)    # #718096 - Muted Grey
    
    # Configure Normal Style
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Segoe UI'
    style_normal.font.size = Pt(10.5)
    style_normal.font.color.rgb = COLOR_TEXT
    style_normal.paragraph_format.line_spacing = 1.25
    style_normal.paragraph_format.space_after = Pt(6)

    def set_cell_background(cell, color_hex):
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
        cell._tc.get_or_add_tcPr().append(shading)

    def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
        tcPr = cell._tc.get_or_add_tcPr()
        tcMar = OxmlElement('w:tcMar')
        for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
            node = OxmlElement(f'w:{m}')
            node.set(qn('w:w'), str(val))
            node.set(qn('w:type'), 'dxa')
            tcMar.append(node)
        tcPr.append(tcMar)

    def make_callout(paragraph, color_hex="F7FAFC", border_color="195402"):
        pPr = paragraph._p.get_or_add_pPr()
        pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:left w:val="single" w:sz="36" w:space="12" w:color="{border_color}"/></w:pBdr>')
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
        pPr.append(pBdr)
        pPr.append(shd)

    def add_custom_heading(text, level, space_before=18, space_after=6):
        h = doc.add_heading(level=level)
        h.paragraph_format.space_before = Pt(space_before)
        h.paragraph_format.space_after = Pt(space_after)
        h.paragraph_format.keep_with_next = True
        run = h.add_run(text)
        run.font.name = 'Segoe UI'
        run.bold = True
        
        if level == 1:
            run.font.size = Pt(18)
            run.font.color.rgb = COLOR_PRIMARY
            # Add a bottom border line using XML for Heading 1
            pPr = h._p.get_or_add_pPr()
            pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="12" w:space="4" w:color="195402"/></w:pBdr>')
            pPr.append(pBdr)
        elif level == 2:
            run.font.size = Pt(14)
            run.font.color.rgb = COLOR_SECONDARY
        elif level == 3:
            run.font.size = Pt(11.5)
            run.font.color.rgb = COLOR_TEXT
        return h

    # ==========================================
    # --- COVER PAGE (MODERN & PREMIUM) ---
    # ==========================================
    # Top primary color accent bar
    accent_table = doc.add_table(rows=1, cols=1)
    accent_table.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
    cell = accent_table.cell(0, 0)
    set_cell_background(cell, "195402")
    set_cell_margins(cell, top=160, bottom=160, left=150, right=150)
    accent_p = cell.paragraphs[0]
    accent_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = accent_p.add_run("KIMERP PLATFORM INTEGRATION SYSTEM")
    run.font.name = 'Segoe UI'
    run.font.size = Pt(11)
    run.bold = True
    run.font.color.rgb = RGBColor(255, 255, 255)

    # Empty vertical space
    for _ in range(5):
        doc.add_paragraph()

    # Main Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_title = title_p.add_run("KimERP System Reference Manual")
    run_title.font.name = 'Segoe UI'
    run_title.font.size = Pt(28)
    run_title.bold = True
    run_title.font.color.rgb = COLOR_PRIMARY

    # Subtitle
    sub_p = doc.add_paragraph()
    sub_p.paragraph_format.space_after = Pt(24)
    run_sub = sub_p.add_run("Developer Onboarding, Integration Guide, Port Configurations, Database Credentials & API Technical Reference")
    run_sub.font.name = 'Segoe UI'
    run_sub.font.size = Pt(12)
    run_sub.font.color.rgb = COLOR_MUTED

    # Bottom border line under header
    rule_table = doc.add_table(rows=1, cols=1)
    rule_cell = rule_table.cell(0, 0)
    set_cell_background(rule_cell, "E2E8F0")
    set_cell_margins(rule_cell, top=20, bottom=20, left=0, right=0)
    rule_cell.paragraphs[0].paragraph_format.space_after = Pt(0)

    for _ in range(8):
        doc.add_paragraph()

    # Metadata Panel (Table format for visual excellence)
    meta_table = doc.add_table(rows=5, cols=2)
    meta_table.autofit = False
    
    meta_labels = [
        ("Prepared For:", "Konwert India Motors Private Limited"),
        ("System Domain:", "https://kimerp.ddns.net"),
        ("Prepared By:", "Antigravity AI Coding Assistant (DeepMind Team)"),
        ("Documentation Version:", "1.0.0 (Production Release)"),
        ("Release Date:", "May 30, 2026")
    ]

    for idx, (label, val) in enumerate(meta_labels):
        row = meta_table.rows[idx]
        cell_lbl, cell_val = row.cells[0], row.cells[1]
        
        cell_lbl.width = Inches(2.2)
        cell_val.width = Inches(4.3)
        
        set_cell_margins(cell_lbl, top=60, bottom=60, left=100, right=100)
        set_cell_margins(cell_val, top=60, bottom=60, left=100, right=100)
        
        p_lbl = cell_lbl.paragraphs[0]
        run_lbl = p_lbl.add_run(label)
        run_lbl.font.name = 'Segoe UI'
        run_lbl.font.size = Pt(9.5)
        run_lbl.bold = True
        run_lbl.font.color.rgb = COLOR_MUTED
        
        p_val = cell_val.paragraphs[0]
        run_val = p_val.add_run(val)
        run_val.font.name = 'Segoe UI'
        run_val.font.size = Pt(9.5)
        if label == "Prepared For:":
            run_val.bold = True
            run_val.font.color.rgb = COLOR_PRIMARY
        else:
            run_val.font.color.rgb = COLOR_TEXT

    # Page Break after Cover Page
    doc.add_page_break()

    # ==========================================
    # --- TABLE OF CONTENTS ---
    # ==========================================
    add_custom_heading("Document Contents", level=1)
    
    toc_p = doc.add_paragraph()
    toc_p.paragraph_format.space_after = Pt(24)
    run_toc_desc = toc_p.add_run("This technical manual contains comprehensive system architecture, database configurations, connection details, and integrations required for maintaining and developing the KimERP software suite.")
    run_toc_desc.font.italic = True
    
    toc_items = [
        ("1. System Architectural Topology & Data Flow", "Page 3"),
        ("2. Containerized Services & Port Configurations", "Page 4"),
        ("3. Comprehensive Tech Stack & Library Matrix", "Page 5"),
        ("4. Database Management & Production Credentials", "Page 6"),
        ("5. Security Secrets, Key Pairs & Token Expirations", "Page 7"),
        ("6. Biometric Machine LAN Punching & Fetch Sync", "Page 8"),
        ("7. Core API Endpoint Reference Directory", "Page 9"),
        ("8. Codebase Anatomy & Folder Tree Mapping", "Page 11"),
        ("9. Build, Deploy & Let's Encrypt Operations Guide", "Page 12")
    ]
    
    for item, page in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        
        run_item = p.add_run(item.ljust(95, '.'))
        run_item.font.name = 'Segoe UI'
        run_item.font.size = Pt(10)
        run_item.font.color.rgb = COLOR_TEXT
        
        run_page = p.add_run(" " + page)
        run_page.font.name = 'Segoe UI'
        run_page.font.size = Pt(10)
        run_page.bold = True
        run_page.font.color.rgb = COLOR_PRIMARY

    doc.add_page_break()

    # ==========================================
    # --- SECTION 1: ARCHITECTURE ---
    # ==========================================
    add_custom_heading("1. System Architectural Topology & Data Flow", level=1)
    
    p = doc.add_paragraph()
    p.add_run("KimERP is designed as a secure, containerized, multi-tier enterprise architecture. The platform operates on a single Linux physical or cloud-based server hosting three major virtual environments orchestrating Nginx (Web Server and SSL termination), FastAPI (Asynchronous Python API server), and PostgreSQL 15 (Relational Database).")
    
    p2 = doc.add_paragraph()
    p2.add_run("Clients (web browsers and Capacitor-based Android mobile apps) authenticate securely using JSON Web Tokens (JWT) through Nginx. The Nginx reverse proxy routes front-end assets directly from the container's static build files, while standard API requests are proxied internally to the backend daemon running on port 8000.")

    # High-level ASCII diagram
    diagram_p = doc.add_paragraph()
    diagram_p.paragraph_format.space_before = Pt(12)
    diagram_p.paragraph_format.space_after = Pt(12)
    make_callout(diagram_p, color_hex="F8FAFC", border_color="195402")
    
    diagram_text = (
        "+-----------------------------------------------------------------------------------+\n"
        "|                         SYSTEM CONNECTIVITY & DATA FLOW DIAGRAM                   |\n"
        "+-----------------------------------------------------------------------------------+\n"
        "|                                                                                   |\n"
        "|  [ Web Browser Client ] <==== HTTPS (Port 443) ====> [ Nginx Container (Frontend) ]|\n"
        "|                                                              ||                   |\n"
        "|  [ Capacitor Mobile   ] <==== HTTPS (Port 443) ==============//                   |\n"
        "|  (App ID: com.kimerp.app)                                    || (Internal Proxy)  |\n"
        "|                                                              \/                   |\n"
        "|                                               [ FastAPI Container (Backend) ]     |\n"
        "|                                               (Running Asynchronously on :8000)   |\n"
        "|                                                              ||                   |\n"
        "|                                                              \/ (Docker DB Link)   |\n"
        "|                                               [ PostgreSQL Container (DB) ]       |\n"
        "|                                               (Database Engine on Port :5432)     |\n"
        "|                                                                                   |\n"
        "|  [ Biometric Machine ] <==== PyZK TCP Port 4370 ====> [ FastAPI Backend Service ] |\n"
        "|  (Local LAN: 192.168.31.4)                    (Cron scheduler pulls logs)         |\n"
        "|                                                                                   |\n"
        "|  [ WebPush Subscriptions ] <==== JSON payload ======> [ VAPID Push Dispatcher ]   |\n"
        "|                                                                                   |\n"
        "+-----------------------------------------------------------------------------------+"
    )
    
    run_diag = diagram_p.add_run(diagram_text)
    run_diag.font.name = 'Consolas'
    run_diag.font.size = Pt(8.5)
    run_diag.font.color.rgb = COLOR_PRIMARY
    run_diag.bold = True

    add_custom_heading("Client Request Resolution Flow", level=2)
    p3 = doc.add_paragraph()
    p3.add_run("1. Client Requests Page: Web request hits ports 80/443 of the physical server (mapped to `erp_frontend`).\n"
               "2. SSL Handshake: Nginx terminates SSL using certificates supplied dynamically via standard Let's Encrypt volumes.\n"
               "3. API Router: Requests prefixed with `/api` are redirected internally to `http://backend:8000` via the internal Docker bridge subnet.\n"
               "4. Asynchronous Threading: FastAPI handles the routed endpoints asynchronously, fetching or writing transaction details to PostgreSQL as necessary.")

    doc.add_page_break()

    # ==========================================
    # --- SECTION 2: PORTS & SERVICES ---
    # ==========================================
    add_custom_heading("2. Containerized Services & Port Configurations", level=1)
    
    p = doc.add_paragraph()
    p.add_run("All software operations within the system are governed by container configurations managed through Docker Compose version 3.8. The services are isolated in an internal network, mapping only strictly required endpoints to the host physical interface.")

    # Ports Table
    table = doc.add_table(rows=4, cols=6)
    table.autofit = False
    
    headers = ["Service Name", "Container Alias", "Image Source", "Internal Port", "External Port", "Exposed Purpose"]
    col_widths = [Inches(1.1), Inches(1.1), Inches(1.5), Inches(0.9), Inches(0.9), Inches(1.5)]
    
    # Format Headers
    hdr_row = table.rows[0]
    for idx, header in enumerate(headers):
        cell = hdr_row.cells[idx]
        cell.width = col_widths[idx]
        set_cell_background(cell, "195402")
        set_cell_margins(cell, top=120, bottom=120, left=100, right=100)
        p_hdr = cell.paragraphs[0]
        run_hdr = p_hdr.add_run(header)
        run_hdr.font.name = 'Segoe UI'
        run_hdr.font.size = Pt(9)
        run_hdr.bold = True
        run_hdr.font.color.rgb = RGBColor(255, 255, 255)

    services_data = [
        ("Nginx Gateway", "erp_frontend", "Local Dockerfile (frontend)", "80, 443", "80, 443", "Web delivery & SSL Gateway"),
        ("FastAPI Core", "erp_backend", "Local Dockerfile (backend)", "8000", "8000", "API engine, Excel/PDF builds"),
        ("PostgreSQL", "erp_db", "postgres:15-alpine", "5432", "5432", "Primary Relational Database")
    ]

    for row_idx, data in enumerate(services_data):
        row = table.rows[row_idx + 1]
        bg_hex = "F7FAFC" if row_idx % 2 == 0 else "FFFFFF"
        for col_idx, item in enumerate(data):
            cell = row.cells[col_idx]
            cell.width = col_widths[col_idx]
            set_cell_background(cell, bg_hex)
            set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
            
            p_cell = cell.paragraphs[0]
            p_cell.paragraph_format.space_after = Pt(0)
            run_cell = p_cell.add_run(item)
            run_cell.font.name = 'Segoe UI'
            run_cell.font.size = Pt(9.5)
            run_cell.font.color.rgb = COLOR_TEXT
            
            if col_idx == 0:
                run_cell.bold = True

    add_custom_heading("Security Advisory regarding External Port Exposure", level=2)
    warn_p = doc.add_paragraph()
    make_callout(warn_p, color_hex="FFF5F5", border_color="FC8181")
    run_warn = warn_p.add_run("IMPORTANT: In production configurations, port 5432 (PostgreSQL) should be blocked at the system firewall level (UFW/Iptables) for external connections, allowing access strictly from trusted administrator IPs. Internal Docker bridge networks protect container-to-container database traffic natively.")
    run_warn.font.size = Pt(9.5)
    run_warn.font.color.rgb = RGBColor(197, 48, 48)

    doc.add_page_break()

    # ==========================================
    # --- SECTION 3: TECH STACK MATRIX ---
    # ==========================================
    add_custom_heading("3. Comprehensive Tech Stack & Library Matrix", level=1)
    
    p = doc.add_paragraph()
    p.add_run("KimERP enforces high performant backend computation coupled with standard micro-service frontend layouts. Below is the full technology stack and matching package locks verified for platform compatibility:")

    add_custom_heading("Backend Technology & Dependencies", level=2)
    
    # Backend Tech table
    be_table = doc.add_table(rows=12, cols=3)
    be_table.autofit = False
    be_widths = [Inches(1.8), Inches(1.2), Inches(3.5)]
    
    # Headers
    hdr = be_table.rows[0]
    for idx, header in enumerate(["Backend Component", "Version", "Direct Purpose / Application context"]):
        cell = hdr.cells[idx]
        cell.width = be_widths[idx]
        set_cell_background(cell, "195402")
        set_cell_margins(cell, top=100, bottom=100, left=100, right=100)
        p_hdr = cell.paragraphs[0]
        run_hdr = p_hdr.add_run(header)
        run_hdr.font.name = 'Segoe UI'
        run_hdr.font.size = Pt(9)
        run_hdr.bold = True
        run_hdr.font.color.rgb = RGBColor(255, 255, 255)

    be_dependencies = [
        ("FastAPI", "0.111.0", "Core REST framework running async routes"),
        ("Uvicorn[standard]", "0.30.0", "Asynchronous ASGI web application server"),
        ("SQLAlchemy", "2.0.30", "Object-Relational Mapping (ORM) data tier"),
        ("Alembic", "1.13.1", "Incremental database schema version control & migrations"),
        ("psycopg2-binary", "2.9.9", "Native PostgreSQL adapter for python database pools"),
        ("python-jose[crypto]", "3.3.0", "JWT credentials signing, decoding and security controls"),
        ("passlib[bcrypt]", "1.7.4", "Standard salted one-way hashing for user passwords"),
        ("openpyxl", "3.1.2", "Excel spreadsheet generation for payroll and CRM exports"),
        ("xhtml2pdf", "0.2.11", "HTML to PDF conversion engine for monthly payslips"),
        ("pyzk", "0.9", "Low level network sync bindings for ZK biometric attendance"),
        ("APScheduler", "3.10.4", "Chronological job dispatcher for automations & leaves auto-approvals")
    ]

    for row_idx, data in enumerate(be_dependencies):
        row = be_table.rows[row_idx + 1]
        bg_hex = "F7FAFC" if row_idx % 2 == 0 else "FFFFFF"
        for col_idx, item in enumerate(data):
            cell = row.cells[col_idx]
            cell.width = be_widths[col_idx]
            set_cell_background(cell, bg_hex)
            set_cell_margins(cell, top=60, bottom=60, left=100, right=100)
            
            p_cell = cell.paragraphs[0]
            p_cell.paragraph_format.space_after = Pt(0)
            run_cell = p_cell.add_run(item)
            run_cell.font.name = 'Segoe UI'
            run_cell.font.size = Pt(9)
            run_cell.font.color.rgb = COLOR_TEXT
            if col_idx == 0:
                run_cell.bold = True

    # Frontend Stack Heading
    add_custom_heading("Frontend Technology & Mobile Wrappers", level=2)
    
    fe_table = doc.add_table(rows=8, cols=3)
    fe_table.autofit = False
    
    # Headers
    hdr = fe_table.rows[0]
    for idx, header in enumerate(["Frontend Component", "Version", "Application Context"]):
        cell = hdr.cells[idx]
        cell.width = be_widths[idx]
        set_cell_background(cell, "238204")
        set_cell_margins(cell, top=100, bottom=100, left=100, right=100)
        p_hdr = cell.paragraphs[0]
        run_hdr = p_hdr.add_run(header)
        run_hdr.font.name = 'Segoe UI'
        run_hdr.font.size = Pt(9)
        run_hdr.bold = True
        run_hdr.font.color.rgb = RGBColor(255, 255, 255)

    fe_dependencies = [
        ("Vite Core", "8.0.4", "Ultra-fast bundling and module hot reloading"),
        ("React", "19.2.4", "Core component-driven user interface framework"),
        ("React Router DOM", "6.28.0", "Declarative browser and capacitor path routing"),
        ("Axios Client", "1.15.0", "Interceptors-driven REST API connection engine"),
        ("Capacitor Core & CLI", "8.3.4", "Cross-platform mobile wrapper and runtime pipeline"),
        ("Lucide React", "0.469.0", "Responsive clean vectors icon engine"),
        ("Recharts", "3.8.1", "High-performance statistics & analytical charts")
    ]

    for row_idx, data in enumerate(fe_dependencies):
        row = fe_table.rows[row_idx + 1]
        bg_hex = "F7FAFC" if row_idx % 2 == 0 else "FFFFFF"
        for col_idx, item in enumerate(data):
            cell = row.cells[col_idx]
            cell.width = be_widths[col_idx]
            set_cell_background(cell, bg_hex)
            set_cell_margins(cell, top=60, bottom=60, left=100, right=100)
            
            p_cell = cell.paragraphs[0]
            p_cell.paragraph_format.space_after = Pt(0)
            run_cell = p_cell.add_run(item)
            run_cell.font.name = 'Segoe UI'
            run_cell.font.size = Pt(9)
            run_cell.font.color.rgb = COLOR_TEXT
            if col_idx == 0:
                run_cell.bold = True

    doc.add_page_break()

    # ==========================================
    # --- SECTION 4: DATABASE CONFIG ---
    # ==========================================
    add_custom_heading("4. Database Management & Production Credentials", level=1)
    
    p = doc.add_paragraph()
    p.add_run("KimERP database uses PostgreSQL version 15 (packaged inside alpine container to reduce deployment size). Below is the comprehensive setup matrix including configurations, default usernames, passwords, internal docker addresses, and superuser master credentials:")

    # DB Details Callout
    db_p = doc.add_paragraph()
    make_callout(db_p, color_hex="F7FAFC", border_color="195402")
    
    db_text = (
        "PostgreSQL Configuration Details:\n"
        "---------------------------------\n"
        "• Container Service Name:   db\n"
        "• DBMS Platform:            PostgreSQL 15 (Alpine Edition)\n"
        "• Host IP (Internal Bridge): db (Resolvable within docker environment)\n"
        "• Host IP (Local External):  localhost (or physical network server IP)\n"
        "• Port Assignment:          5432\n"
        "• Target Database Name:     erp_db\n"
        "• Database User Name:       erp_user\n"
        "• Database Password:       erp_pass\n"
        "• Timezone Setting:         Asia/Kolkata (Force synchronized with server TZ)\n"
        "• Postgres Data Volume:     postgres_data (Persistent docker named volume)"
    )
    run_db = db_p.add_run(db_text)
    run_db.font.name = 'Consolas'
    run_db.font.size = Pt(9.5)
    run_db.font.color.rgb = COLOR_PRIMARY
    run_db.bold = True

    add_custom_heading("Unified Connection Strings", level=2)
    p_conn = doc.add_paragraph()
    p_conn.add_run("The database backend abstracts pool connections utilizing SQLAlchemy. The standard URLs to reach the data layer in development vs deployment are:")
    
    # Code blocks for connection strings
    p_conn_int = doc.add_paragraph()
    make_callout(p_conn_int, color_hex="F1F5F9", border_color="238204")
    run_int = p_conn_int.add_run(
        "# Internal Docker Connection string (Backend Service to Database Service)\n"
        "DATABASE_URL=postgresql://erp_user:erp_pass@db:5432/erp_db"
    )
    run_int.font.name = 'Consolas'
    run_int.font.size = Pt(9)
    run_int.font.color.rgb = COLOR_TEXT
    
    p_conn_ext = doc.add_paragraph()
    make_callout(p_conn_ext, color_hex="F1F5F9", border_color="238204")
    run_ext = p_conn_ext.add_run(
        "# External Host Connection string (Direct development connection or administration scripts)\n"
        "DATABASE_URL=postgresql://erp_user:erp_pass@localhost:5432/erp_db"
    )
    run_ext.font.name = 'Consolas'
    run_ext.font.size = Pt(9)
    run_ext.font.color.rgb = COLOR_TEXT

    add_custom_heading("Initial Seeding & Default Admin Master Credentials", level=2)
    p_seed = doc.add_paragraph()
    p_seed.add_run("Upon initializing a fresh instance of the system via docker orchestration, standard admin accounts, branch default values, department hierarchies, and customizable form structures are seeded. The master super admin credential seeded by default is:")

    # Admin info Callout
    admin_p = doc.add_paragraph()
    make_callout(admin_p, color_hex="FFFDF5", border_color="B7791F") # Amber accent
    admin_text = (
        "Default Master Administrator Credentials:\n"
        "-----------------------------------------\n"
        "• Administrator Login User Email: admin@erp.com\n"
        "• Seeded Super Admin Password:    admin123\n"
        "• Role Assignment:               SuperAdmin (All permissions bypassed)"
    )
    run_admin = admin_p.add_run(admin_text)
    run_admin.font.name = 'Consolas'
    run_admin.font.size = Pt(9.5)
    run_admin.bold = True
    run_admin.font.color.rgb = RGBColor(183, 121, 31)

    doc.add_page_break()

    # ==========================================
    # --- SECTION 5: APP SECRETS ---
    # ==========================================
    add_custom_heading("5. Security Secrets, Key Pairs & Token Expirations", level=1)
    
    p = doc.add_paragraph()
    p.add_run("KimERP enforces high cryptographic constraints to prevent session-jacking, cross-site scripting (XSS), or unauthorized payload injection. System configurations rely strictly on environment variables read directly by Pydantic settings at execution.")

    add_custom_heading("JWT Authentication Tokens Config", level=2)
    p_jwt = doc.add_paragraph()
    p_jwt.add_run("Tokens are signed on successful authentication and validated at each restricted route header. Key credentials stored natively in the application core settings are:")

    jwt_tbl = doc.add_table(rows=4, cols=2)
    jwt_widths = [Inches(2.5), Inches(4.0)]
    
    jwt_data = [
        ("SECRET_KEY", "omni-erp-secret-2024 (Can be overridden in .env at deployment)"),
        ("ALGORITHM", "HS256 (HMAC SHA-256 Symmetric encryption)"),
        ("ACCESS_TOKEN_EXPIRE_MINUTES", "480 Minutes (Equivalent to an 8-Hour standard working shift)")
    ]

    for idx, (param, val) in enumerate(jwt_data):
        row = jwt_tbl.rows[idx]
        set_cell_background(row.cells[0], "F7FAFC")
        set_cell_background(row.cells[1], "FFFFFF")
        
        set_cell_margins(row.cells[0], top=80, bottom=80, left=100, right=100)
        set_cell_margins(row.cells[1], top=80, bottom=80, left=100, right=100)
        
        row.cells[0].width = jwt_widths[0]
        row.cells[1].width = jwt_widths[1]
        
        p_c1 = row.cells[0].paragraphs[0]
        run_c1 = p_c1.add_run(param)
        run_c1.font.name = 'Consolas'
        run_c1.font.size = Pt(9.5)
        run_c1.bold = True
        run_c1.font.color.rgb = COLOR_PRIMARY
        
        p_c2 = row.cells[1].paragraphs[0]
        run_c2 = p_c2.add_run(val)
        run_c2.font.name = 'Segoe UI'
        run_c2.font.size = Pt(9.5)
        run_c2.font.color.rgb = COLOR_TEXT

    add_custom_heading("Push Notifications Web VAPID Keys Reference", level=2)
    p_vapid = doc.add_paragraph()
    p_vapid.add_run("The application backend natively supports real-time Web Push notifications (for attendance reminders, leave approvals, shifts alerts) routed securely over standard browser VAPID connections. The configuration key pairs used internally are:")

    # Vapid PEM Callout
    vapid_p = doc.add_paragraph()
    make_callout(vapid_p, color_hex="F8FAFC", border_color="195402")
    
    vapid_text = (
        "VAPID Cryptographic Profile Configuration:\n"
        "------------------------------------------\n"
        "• VAPID Contact Email Address: mailto:rkyuvaa@gmail.com\n\n"
        "• VAPID Public Key (For client subscription binding):\n"
        "  BEHYGaYuo_Ye0vvnjvqldoHRx83a8Tn_300kiwcGsIvE9was0CbXkAooFJVMcf4FTg9UXEj6ww5ndEjje5F0X_U\n\n"
        "• VAPID Private Key (RSA Private PEM - DO NOT SHARE EXTERNALLY):\n"
        "  -----BEGIN PRIVATE KEY-----\n"
        "  MIGHAgEAMBMGByqGSM49AgEGCCqGSM49AwEHBG0wawIBAQQgDLpF3hVCR+bWv5bg\n"
        "  TuEraBQZFYCSxOaKwJ7d8eWivlGhRANCAARB2BmmLqP2HtL75476pXaB0cfN2vE5\n"
        "  /99NJIsHBrCLxPcGrNAm15AKKBSVTHH+BU4PVFxI+sMOZ3RI43uRdF/1\n"
        "  -----END PRIVATE KEY-----"
    )
    run_vapid_txt = vapid_p.add_run(vapid_text)
    run_vapid_txt.font.name = 'Consolas'
    run_vapid_txt.font.size = Pt(8.5)
    run_vapid_txt.font.color.rgb = COLOR_PRIMARY
    run_vapid_txt.bold = True

    doc.add_page_break()

    # ==========================================
    # --- SECTION 6: BIOMETRIC SYSTEMS ---
    # ==========================================
    add_custom_heading("6. Biometric Machine LAN Punching & Fetch Sync", level=1)
    
    p = doc.add_paragraph()
    p.add_run("KimERP hosts a seamless, native integration with ZK Teco physical biometric devices. The platform bridges local biometric attendance machines with the cloud FastAPI backend dynamically.")

    add_custom_heading("Local Connection Architecture", level=2)
    p_bio = doc.add_paragraph()
    p_bio.add_run("1. Network Port Binding: Biometric physical machines operate by default on UDP/TCP Port 4370.\n"
               "2. Local IP Punching: Devices reside in the company's local network subnets (typically mapped to static IPs, e.g. 192.168.31.4).\n"
               "3. Asynchronous Cron Fetch: An integrated scheduler in the backend queries machine endpoints every 10 minutes to fetch punch streams and reconcile with employee database rosters.")

    # PyZK details Callout
    pyzk_p = doc.add_paragraph()
    make_callout(pyzk_p, color_hex="F7FAFC", border_color="238204")
    
    pyzk_text = (
        "Biometric Cron Synchronizer Logic:\n"
        "-----------------------------------\n"
        "• Core Python library:        pyzk (Version 0.9)\n"
        "• Dispatch Schedule:          Every 10 minutes (Triggered via APScheduler)\n"
        "• Active Scheduler Flag:      DISABLE_BIOMETRIC_SCHEDULER (Configurable via Environment)\n"
        "• Fetch Endpoint:             sync_all_biometric_machines()\n"
        "• Sync Execution Process:     _do_sync(machine_id)\n"
        "• Auto-Approval Window:       Leaves/On-Duty requests auto-approve if pending for 6 hours"
    )
    run_pyzk = pyzk_p.add_run(pyzk_text)
    run_pyzk.font.name = 'Consolas'
    run_pyzk.font.size = Pt(9.5)
    run_pyzk.bold = True
    run_pyzk.font.color.rgb = COLOR_PRIMARY

    add_custom_heading("Biometric Cron Control Environment Flags", level=2)
    p_flags = doc.add_paragraph()
    p_flags.add_run("To bypass or disable automatic background polling (for instance, during development runs or system debugging to prevent socket errors on inaccessible LAN endpoints), developers can manage the schedule using environment variables:")
    
    p_code = doc.add_paragraph()
    make_callout(p_code, color_hex="F1F5F9", border_color="195402")
    run_code = p_code.add_run(
        "# Set inside .env file or server launch configuration to silence automatic biometric queries\n"
        "DISABLE_BIOMETRIC_SCHEDULER=True"
    )
    run_code.font.name = 'Consolas'
    run_code.font.size = Pt(9)
    run_code.font.color.rgb = COLOR_TEXT

    doc.add_page_break()

    # ==========================================
    # --- SECTION 7: API REFERENCE ---
    # ==========================================
    add_custom_heading("7. Core API Endpoint Reference Directory", level=1)
    
    p = doc.add_paragraph()
    p.add_run("All operations are built around a unified API routing structure. The base routing prefix is '/api'. Below is the directory mapping of core endpoints for developers:")

    api_widths = [Inches(2.5), Inches(1.0), Inches(3.0)]

    def create_api_table(rows):
        t = doc.add_table(rows=rows, cols=3)
        t.autofit = False
        hdr = t.rows[0]
        for idx, h in enumerate(["API Endpoint Router", "Methods", "Functional Scope"]):
            cell = hdr.cells[idx]
            cell.width = api_widths[idx]
            set_cell_background(cell, "195402")
            set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
            p_h = cell.paragraphs[0]
            r_h = p_h.add_run(h)
            r_h.font.name = 'Segoe UI'
            r_h.font.size = Pt(9)
            r_h.bold = True
            r_h.font.color.rgb = RGBColor(255, 255, 255)
        return t

    def fill_api_row(table, row_idx, ep, method, desc, bg_hex):
        row = table.rows[row_idx]
        for col_idx, item in enumerate([ep, method, desc]):
            cell = row.cells[col_idx]
            cell.width = api_widths[col_idx]
            set_cell_background(cell, bg_hex)
            set_cell_margins(cell, top=60, bottom=60, left=100, right=100)
            p_cell = cell.paragraphs[0]
            p_cell.paragraph_format.space_after = Pt(0)
            
            run_c = p_cell.add_run(item)
            run_c.font.name = 'Consolas' if col_idx < 2 else 'Segoe UI'
            run_c.font.size = Pt(8.5) if col_idx < 2 else Pt(9)
            run_c.font.color.rgb = COLOR_TEXT
            if col_idx == 0:
                run_c.bold = True
                run_c.font.color.rgb = COLOR_PRIMARY

    # Auth & Users Table
    add_custom_heading("Authentication & Security Module Router", level=2)
    auth_endpoints = [
        ("/api/auth/login", "POST", "User login, validates credentials and issues JWT access token"),
        ("/api/auth/me", "GET", "Retrieves active authenticated profile metadata"),
        ("/api/users/", "GET/POST", "Fetches or registers primary system users"),
        ("/api/users/{id}", "PUT/DELETE", "Updates or revokes specific user accounts"),
        ("/api/roles/", "GET/POST", "Manages user system permissions (SuperAdmin, Admin, Employee, Partner)")
    ]
    t_auth = create_api_table(len(auth_endpoints) + 1)
    for idx, (ep, method, desc) in enumerate(auth_endpoints):
        bg = "F7FAFC" if idx % 2 == 0 else "FFFFFF"
        fill_api_row(t_auth, idx + 1, ep, method, desc, bg)

    # HR & Attendance Table
    add_custom_heading("HR, Leave Management & Biometric Router", level=2)
    hr_endpoints = [
        ("/api/hr/employees/", "GET/POST", "Employee master record database, manages salaries & attachments"),
        ("/api/hr/shifts/", "GET/POST", "Configures structural shift rules (grace timings, hours, overtime)"),
        ("/api/hr/leave/", "GET/POST", "Submits, reviews or handles personal leave requests & balances"),
        ("/api/hr/attendance/punch", "POST", "API backend punch endpoint (supports selfie uploads and GPS checks)"),
        ("/api/hr/attendance/logs", "GET", "Renders system logs and monthly employee sheet summaries"),
        ("/api/hr/biometric/", "GET/POST", "Configures, tests and coordinates local biometric connections"),
        ("/api/hr/payroll/generate", "POST", "Runs standard salary formulas and creates monthly payroll runs"),
        ("/api/hr/payroll/payslip/{id}", "GET", "Compiles dynamic HTML to PDF and dispatches payslip copies")
    ]
    t_hr = create_api_table(len(hr_endpoints) + 1)
    for idx, (ep, method, desc) in enumerate(hr_endpoints):
        bg = "F7FAFC" if idx % 2 == 0 else "FFFFFF"
        fill_api_row(t_hr, idx + 1, ep, method, desc, bg)

    # Studio & Form Engine Table
    add_custom_heading("Dynamic Forms Studio & Operation Modules Router", level=2)
    op_endpoints = [
        ("/api/studio/fields", "POST/PUT", "Drags, drops or mutates customized layout configurations"),
        ("/api/forms/{module_name}", "GET/POST", "Saves or fetches values for customized form frameworks"),
        ("/api/crm/", "GET/POST", "Handles primary customer leads and lifecycle activities"),
        ("/api/installation/", "GET/POST", "Coordinates installation planning and field deployments"),
        ("/api/service/", "GET/POST", "Administers customer issue logs and technician dispatches"),
        ("/api/warranty/", "GET/POST", "Registers, updates and logs active product warranties")
    ]
    t_op = create_api_table(len(op_endpoints) + 1)
    for idx, (ep, method, desc) in enumerate(op_endpoints):
        bg = "F7FAFC" if idx % 2 == 0 else "FFFFFF"
        fill_api_row(t_op, idx + 1, ep, method, desc, bg)

    doc.add_page_break()

    # ==========================================
    # --- SECTION 8: FILE STRUCTURE ---
    # ==========================================
    add_custom_heading("8. Codebase Anatomy & Folder Tree Mapping", level=1)
    
    p = doc.add_paragraph()
    p.add_run("For seamless onboarding, developers should familiarize themselves with the core system folder tree. Below is the file mapping indicating backend routing, frontend components, and orchestration scripts:")

    # File Structure Callout
    fs_p = doc.add_paragraph()
    make_callout(fs_p, color_hex="F8FAFC", border_color="195402")
    
    fs_text = (
        "KimERP Root Project Folder Map:\n"
        "-------------------------------\n"
        "erp/\n"
        "├── docker-compose.yml           # Multi-container service definitions\n"
        "├── .env.example                 # Template for system deployment configuration\n"
        "├── nginx.conf                   # Nginx reverse proxy SSL & API path routing\n"
        "├── backend/                     # Python FastAPI Service Directory\n"
        "│   ├── requirements.txt         # Package dependencies locks\n"
        "│   ├── Dockerfile               # Backend alpine container setup\n"
        "│   ├── backup_manager.py        # Database backup scheduler script\n"
        "│   ├── main.py                  # Entry Point - mounts routers, DB check & APScheduler\n"
        "│   └── app/\n"
        "│       ├── config.py            # App configuration setting class (Pydantic Settings)\n"
        "│       ├── database.py          # SQLAlchemy engine, session pools & Base\n"
        "│       ├── models.py            # Core Database schemas (User, CRM, Forms)\n"
        "│       ├── hr_models.py         # HR Module Database schemas (Employee, Punch, Payroll)\n"
        "│       ├── hr_scheduler.py      # Background jobs (Biometric Sync, Leave Auto-approvals)\n"
        "│       └── routers/             # Independent modular endpoint handlers\n"
        "└── frontend/                    # Vite React Single Page Application (SPA)\n"
        "    ├── package.json             # Bundle scripts and dependencies\n"
        "    ├── capacitor.config.json    # Capacitor Android application configuration\n"
        "    ├── vite.config.js           # Compilation config\n"
        "    ├── android/                 # Auto-generated capacitor native android environment\n"
        "    └── src/\n"
        "        ├── index.css            # Branding design tokens (Deep Green Theme --accent)\n"
        "        ├── main.jsx             # React framework entry\n"
        "        ├── App.jsx              # Main routing shell\n"
        "        └── utils/\n"
        "            ├── api.js           # Axios base request wrapper & routing logic\n"
        "            └── pushNotifications.js # WebPush service integrations"
    )
    run_fs = fs_p.add_run(fs_text)
    run_fs.font.name = 'Consolas'
    run_fs.font.size = Pt(8.5)
    run_fs.font.color.rgb = COLOR_TEXT

    add_custom_heading("Static & File Upload Locations", level=2)
    p_upload = doc.add_paragraph()
    p_upload.add_run("Any uploads (like employee profile pictures, selfie punches, dynamic form documents, CRM contract attachments) are written dynamically to standard physical folders. These directories are mounted persistently into backend docker subdirectories:")
    
    # Upload folder Callout
    folder_p = doc.add_paragraph()
    make_callout(folder_p, color_hex="F1F5F9", border_color="238204")
    folder_text = (
        "Mounted Volume Upload Directories:\n"
        "-----------------------------------\n"
        "• Root Upload Folder:        /app/static/uploads (Internal Container Path)\n"
        "• Host Persistent Volume:    backend_uploads (Docker named volume)\n"
        "• Employee Photo uploads:    /app/static/uploads (Served via /api/uploads/{name})\n"
        "• Selfie Punch uploads:      /app/static/uploads/attendance (Served via /api/uploads/attendance/{name})"
    )
    run_folder = folder_p.add_run(folder_text)
    run_folder.font.name = 'Consolas'
    run_folder.font.size = Pt(9.5)
    run_folder.font.color.rgb = COLOR_TEXT

    doc.add_page_break()

    # ==========================================
    # --- SECTION 9: DEPLOYMENT GUIDE ---
    # ==========================================
    add_custom_heading("9. Build, Deploy & Let's Encrypt Operations Guide", level=1)
    
    p = doc.add_paragraph()
    p.add_run("KimERP is engineered for rapid continuous delivery. A new environment can be initialized on clean servers by executing standard commands outlined below:")

    add_custom_heading("Primary Build & Startup Orchestration", level=2)
    p_build = doc.add_paragraph()
    p_build.add_run("1. Environment Variables: Clone the template `.env.example` file to `.env` and fill the variables.\n"
               "2. Compile and Launch: Run Docker Compose command synchronously in the root project directory:")

    # Build code block
    p_code1 = doc.add_paragraph()
    make_callout(p_code1, color_hex="F1F5F9", border_color="195402")
    run_b1 = p_code1.add_run(
        "# Compile containers, setup persistent volumes, initialize internal subnets and launch in detached mode\n"
        "docker-compose up --build -d\n\n"
        "# Verify that all three core containers are operating correctly\n"
        "docker-compose ps"
    )
    run_b1.font.name = 'Consolas'
    run_b1.font.size = Pt(9)
    run_b1.font.color.rgb = COLOR_TEXT

    add_custom_heading("Dynamic SSL Handshake Configuration", level=2)
    p_ssl = doc.add_paragraph()
    p_ssl.add_run("SSL termination is handled by `erp_frontend` container via Nginx. During initialization, Let's Encrypt certificates from the physical host are mapped inside the nginx volume space in read-only access format:")
    
    p_code2 = doc.add_paragraph()
    make_callout(p_code2, color_hex="F1F5F9", border_color="195402")
    run_b2 = p_code2.add_run(
        "# Docker Compose Nginx SSL volume mount reference:\n"
        "volumes:\n"
        "  - /etc/letsencrypt:/etc/letsencrypt:ro"
    )
    run_b2.font.name = 'Consolas'
    run_b2.font.size = Pt(9)
    run_b2.font.color.rgb = COLOR_TEXT

    add_custom_heading("Mobile Capacitor Native Compilation Pipeline", level=2)
    p_cap = doc.add_paragraph()
    p_cap.add_run("The native mobile wrapper is maintained using Ionic Capacitor. To compile backend API bindings into a deployable Android Package (APK):")

    p_code3 = doc.add_paragraph()
    make_callout(p_code3, color_hex="F1F5F9", border_color="195402")
    run_b3 = p_code3.add_run(
        "# 1. Compile the React codebase\n"
        "cd frontend\n"
        "npm run build\n\n"
        "# 2. Sync public assets into Capacitor's Android folder mapping\n"
        "npx cap sync\n\n"
        "# 3. Open Android Studio to build the production signed release APK\n"
        "npx cap open android"
    )
    run_b3.font.name = 'Consolas'
    run_b3.font.size = Pt(9)
    run_b3.font.color.rgb = COLOR_TEXT

    # --- END NOTE / FOOTER SIGN OFF ---
    for _ in range(3):
        doc.add_paragraph()
    
    p_end = doc.add_paragraph()
    p_end.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_end = p_end.add_run("--- End of Technical Reference Manual ---")
    run_end.font.name = 'Segoe UI'
    run_end.font.size = Pt(10)
    run_end.font.color.rgb = COLOR_MUTED
    run_end.font.italic = True

    # Save Document
    target_path = r"C:\Users\rkyuv\OneDrive\Documents\erp\KimERP_Developer_Reference.docx"
    doc.save(target_path)
    print(f"SUCCESS: Document generated and saved to {target_path}")

if __name__ == "__main__":
    create_document()
